from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "final_gait_model_method_summary.docx"


BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 16 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9)
        set_cell_fill(table.rows[0].cells[i], GRAY_FILL)
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.color.rgb = BLUE if level <= 2 else DARK


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("최종 보행 평가 모델 방법 요약")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(20)
    run.font.color.rgb = BLUE
    run.bold = True

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Axis-aligned + PhysioNet normal-reference domain correction pipeline")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(80, 80, 80)

    add_heading(doc, "1. 최종 결론", 1)
    p = doc.add_paragraph()
    p.add_run("최종 Flask 서비스에 연결된 보행 모델은 ").bold = False
    r = p.add_run("PhysioNet LabWalks 정상군 기준 보정 모델")
    r.bold = True
    r.font.color.rgb = DARK
    p.add_run("이다. APK CSV를 업로드하면 서버에서 축정렬, 100 Hz 보간, 10초 best window 추출, 최종 feature 계산, 모델 추론을 수행한다.")

    add_bullet(doc, "최종 model artifact: MOCA/models/gait_axis_aligned_physionet_youden.joblib")
    add_bullet(doc, "최종 service extractor: MOCA/gait_axis_aligned_processor.py")
    add_bullet(doc, "Flask 연결 route: MOCA/app.py 의 /gait/upload-csv")
    add_bullet(doc, "분석/학습 스크립트: analysis_scripts/train_final_axis_aligned_domain_corrected_gait_model.py")

    add_heading(doc, "2. 전처리 및 보정 기준", 1)
    add_table(
        doc,
        ["단계", "적용 기준", "목적"],
        [
            ["CSV 입력", "Timestamp_ns, Acc_Vertical_g/ML/AP 또는 Acc_X/Y/Z", "APK 측정값 수신"],
            ["축정렬", "V/ML/AP가 있으면 그대로 사용, raw축은 gravity로 vertical 추정", "기기 방향 차이 감소"],
            ["보간", "100 Hz uniform resampling", "PhysioNet 기준 sampling rate와 정렬"],
            ["필터", "0.6-3 Hz gait band", "보행 리듬 대역만 사용"],
            ["window", "20초 측정 내 10초 best quality window", "가장 안정적인 보행 구간 선택"],
            ["도메인 보정", "PhysioNet LabWalks 정상군 median 기준", "외부 데이터셋별 정상군 위치 차이 보정"],
        ],
        widths=[1.2, 2.7, 2.3],
    )
    doc.add_paragraph(
        "도메인 보정은 정상군 기준으로만 delta를 계산하고, 같은 데이터셋의 정상/저하군 모두에 동일한 delta를 적용했다. "
        "OUR_SAMPLE은 보정값 계산과 학습에 사용하지 않고 홀드아웃 확인용으로만 사용했다."
    )

    add_heading(doc, "3. 최종 사용 feature 정의", 1)
    add_table(
        doc,
        ["Feature", "축", "계산 정의", "해석"],
        [
            ["v_acf_stride_peak", "Vertical", "수직 가속도 ACF에서 stride lag 근처 peak 높이", "수직 보행 반복성/stride regularity"],
            ["v_acf_stride_peak_width_sec", "Vertical", "수직 ACF stride peak의 half-height width", "stride 반복성의 시간적 퍼짐"],
            ["ap_acf_stride_peak_width_sec", "AP", "앞뒤축 ACF stride peak의 half-height width", "앞뒤 방향 stride timing 안정성"],
            ["ap_spec_entropy", "AP", "0.6-3 Hz band power 분포의 normalized spectral entropy", "앞뒤 보행 리듬의 복잡도/분산"],
        ],
        widths=[1.55, 0.75, 2.25, 1.75],
    )

    add_heading(doc, "4. 논문 기반 근거", 1)
    add_bullet(doc, "Autocorrelation 기반 step/stride regularity는 trunk 또는 waist-mounted accelerometer 보행 분석에서 반복적으로 사용되는 방식이다.")
    add_bullet(doc, "Vertical, mediolateral, anteroposterior 축별 acceleration pattern의 regularity/symmetry는 노화 및 보행 기능 차이를 설명하는 데 사용되어 왔다.")
    add_bullet(doc, "Spectral entropy 및 주파수 기반 feature는 보행 리듬의 복잡도, 예측 가능성, freezing/gait impairment 탐지 계열 연구에서 쓰이는 신호 특성이다.")
    add_table(
        doc,
        ["근거", "모델 반영"],
        [
            ["Moe-Nilssen & Helbostad, 2004: trunk accelerometry와 unbiased autocorrelation으로 gait regularity/symmetry 평가", "v_acf_stride_peak, ACF peak width 계열"],
            ["Kobsar et al., 2014: waist-mounted tri-axial accelerometer로 vertical/ML/AP regularity와 symmetry 비교", "축별 ACF feature 사용 근거"],
            ["Scalera et al., 2020: IMU 기반 gait regularity를 autocorrelation으로 산출", "wearable IMU regularity 산출 근거"],
            ["Mazilu et al., 2013 및 PD entropy 연구: entropy/complexity feature로 보행 장애 특성 평가", "ap_spec_entropy 사용 근거"],
        ],
        widths=[3.2, 3.0],
    )

    add_heading(doc, "5. 데이터 구성 및 N수", 1)
    add_table(
        doc,
        ["Dataset", "라벨", "N", "사용 목적"],
        [
            ["PhysioNet LabWalks", "정상", "38", "정상 reference 및 학습"],
            ["PhysioNet LabWalks", "저하", "35", "학습"],
            ["UCI HAR walking", "정상", "30", "외부 스마트폰 정상 보강"],
            ["GEOTEC smartphone walking/TUG", "정상", "10", "외부 스마트폰 정상 보강"],
            ["Chapman PD OFF raw walking", "저하", "20", "운동저하/PD 보행 보강"],
            ["FoG-STAR back walking", "저하", "67", "운동저하/PD 보행 보강"],
            ["OUR_SAMPLE APK", "정상 2, 저하 1", "3", "홀드아웃 smoke test"],
        ],
        widths=[2.3, 1.0, 0.55, 2.35],
    )

    add_heading(doc, "6. 검증 결과", 1)
    add_table(
        doc,
        ["검증", "AUC", "Acc", "Sensitivity", "Specificity", "F1", "Threshold"],
        [
            ["5-fold GroupKFold OOF", "0.851", "0.810", "0.934", "0.615", "0.857", "0.28"],
            ["최종 public train Youden", "0.870", "0.750", "0.648", "0.910", "0.760", "0.56"],
        ],
        widths=[1.8, 0.7, 0.7, 0.9, 0.9, 0.7, 0.9],
    )
    doc.add_paragraph(
        "서비스 배포 threshold는 최종 public 학습셋 Youden 기준 0.56으로 저장했다. "
        "OOF threshold 0.28은 교차검증 검증값으로 metadata에 함께 남겼다."
    )

    add_table(
        doc,
        ["APK SAMPLE", "Target", "Probability", "Threshold", "Prediction", "Correct"],
        [
            ["20260715_163129", "0", "0.378", "0.56", "0", "True"],
            ["20260716_발다침_좌회전", "1", "0.934", "0.56", "1", "True"],
            ["calibrated_155029", "0", "0.213", "0.56", "0", "True"],
        ],
        widths=[2.25, 0.65, 0.9, 0.8, 0.8, 0.7],
    )

    add_heading(doc, "7. 폴더 및 파일 정리", 1)
    add_table(
        doc,
        ["구분", "파일/폴더", "역할"],
        [
            ["Flask runtime", "MOCA/app.py", "CSV 업로드 route와 결과 저장"],
            ["Flask runtime", "MOCA/gait_axis_aligned_processor.py", "최종 CSV 전처리/extractor/predictor"],
            ["Flask runtime", "MOCA/models/gait_axis_aligned_physionet_youden.joblib", "최종 모델"],
            ["Flask runtime", "MOCA/models/gait_axis_aligned_physionet_youden_metadata.json", "최종 모델 metadata"],
            ["Analysis", "analysis_scripts/build_axis_aligned_gait_dataset_and_model.py", "축정렬 subject table 생성"],
            ["Analysis", "analysis_scripts/model_axis_aligned_domain_corrected_gait.py", "도메인 보정 후보 탐색"],
            ["Analysis", "analysis_scripts/train_final_axis_aligned_domain_corrected_gait_model.py", "최종 모델 학습/export"],
            ["Manifest", "docs/final_service_manifest/", "최종/분석 파일 구분 설명"],
        ],
        widths=[1.15, 2.75, 2.3],
    )
    doc.add_paragraph(
        "현재 서비스 extractor와 분석 extractor는 같은 기준으로 구현되어 있으며, SAMPLE smoke test에서 동일한 판정을 냈다. "
        "다만 runtime 안정성을 위해 Flask는 MOCA/gait_axis_aligned_processor.py를 직접 사용하고, 분석 스크립트는 analysis_scripts에 분리해 두었다."
    )

    add_heading(doc, "8. 참고 문헌", 1)
    refs = [
        "Moe-Nilssen R, Helbostad JL. Estimation of gait cycle characteristics by trunk accelerometry. Journal of Biomechanics. 2004. https://doi.org/10.1016/S0021-9290(03)00233-1",
        "Kobsar D et al. Evaluation of age-related differences in the stride-to-stride fluctuations, regularity and symmetry of gait using a waist-mounted tri-axial accelerometer. Gait & Posture. 2014. https://pubmed.ncbi.nlm.nih.gov/24139685/",
        "Scalera GM et al. Gait regularity assessed by wearable sensors: comparison between accelerometer and gyroscope data. Journal of Biomechanics. 2020. https://doi.org/10.1016/j.jbiomech.2020.110115",
        "Mazilu S et al. Automatic detection of freezing of gait events in patients with Parkinson's disease. Computer Methods and Programs in Biomedicine. 2013. https://doi.org/10.1016/j.cmpb.2012.10.016",
        "Real-time gait cycle parameter recognition using a wearable accelerometry system. Sensors. 2011. https://pmc.ncbi.nlm.nih.gov/articles/PMC3231731/",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

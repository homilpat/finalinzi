from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


SRC = Path(r"C:\Users\whdgu\Desktop\파이널 프로젝트\final_보행설명_nestedCV_최종.docx")
DOCX = Path(r"C:\Users\whdgu\Desktop\파이널 프로젝트\final_보행설명_nestedCV_최종_CI포함.docx")

doc = Document(SRC)

auc_ci_text = "AUC 0.830 (95% bootstrap CI: 0.720-0.922)"

doc.paragraphs[55].text = (
    "* 대표 성능은 A scheme(5-fold x 100 repeats) subject-level pooled OOF 기준으로 "
    "AUC 0.830 (95% bootstrap CI: 0.720-0.922), sensitivity 0.800, specificity 0.738이다."
)

# Table 8 and Table 10 A scheme AUC cell.
doc.tables[8].rows[1].cells[1].text = "0.830 (95% CI 0.720-0.922)"
doc.tables[10].rows[1].cells[1].text = "0.830 (95% CI 0.720-0.922)"
doc.tables[11].rows[5].cells[2].text = "0.830 (95% CI 0.720-0.922; A pooled OOF)"

marker = "운동저하 25명 중 20명 정탐"
for i, paragraph in enumerate(doc.paragraphs):
    if marker in paragraph.text:
        target = paragraph
        break
else:
    target = None

if target is not None:
    new_p = OxmlElement("w:p")
    target._p.addnext(new_p)
    inserted = doc.paragraphs[i + 1]
    inserted.text = (
        "AUC 신뢰구간은 A scheme subject-level pooled OOF 평균 확률을 기준으로, "
        "67명을 subject 단위로 10,000회 bootstrap resampling하여 산출하였다. "
        "N=67(positive=25, negative=42)이므로 점추정 AUC와 함께 CI를 병기해 일반화 불확실성을 명시한다."
    )

doc.save(DOCX)
print(DOCX)

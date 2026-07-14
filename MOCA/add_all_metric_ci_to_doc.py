from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement


BASE = Path(r"C:\Users\whdgu\Desktop\파이널 프로젝트")
PRED = BASE / "final__2026" / "02_model" / "domain4_nested_oof_predictions.csv"
CI_CSV = BASE / "final__2026" / "02_model" / "domain4_nested_metric_bootstrap_ci.csv"
SRC = BASE / "final_보행설명_nestedCV_최종_CI포함.docx"
OUT = BASE / "final_보행설명_nestedCV_최종_전체CI포함.docx"


def auc_score(y, s):
    y = np.asarray(y)
    s = np.asarray(s)
    pos = s[y == 1]
    neg = s[y == 0]
    if len(pos) == 0 or len(neg) == 0:
        return np.nan
    wins = (pos[:, None] > neg[None, :]).sum()
    ties = (pos[:, None] == neg[None, :]).sum()
    return float((wins + 0.5 * ties) / (len(pos) * len(neg)))


def metric_values(y, prob, threshold):
    pred = (prob >= threshold).astype(int)
    tp = int(((y == 1) & (pred == 1)).sum())
    fn = int(((y == 1) & (pred == 0)).sum())
    tn = int(((y == 0) & (pred == 0)).sum())
    fp = int(((y == 0) & (pred == 1)).sum())
    return {
        "auc": auc_score(y, prob),
        "accuracy": (tp + tn) / len(y),
        "sensitivity": tp / (tp + fn) if (tp + fn) else np.nan,
        "specificity": tn / (tn + fp) if (tn + fp) else np.nan,
        "f1": 2 * tp / (2 * tp + fp + fn) if (2 * tp + fp + fn) else np.nan,
        "tn": tn,
        "fp": fp,
        "fn": fn,
        "tp": tp,
    }


def fmt(metric, rows):
    row = rows[metric]
    return f"{row['estimate']:.3f} (95% CI {row['ci_low']:.3f}-{row['ci_high']:.3f})"


df = pd.read_csv(PRED)
a = df[(df["scheme"] == "A_5fold_x100") & (df["threshold_strategy"] == "nested_inner_oof_youden")]
pooled = (
    a.groupby("subject_id", as_index=False)
    .agg(target=("target", "first"), probability=("probability", "mean"), threshold=("threshold", "mean"))
)
y = pooled["target"].to_numpy()
prob = pooled["probability"].to_numpy()
threshold = pooled["threshold"].to_numpy()

point = metric_values(y, prob, threshold)
rng = np.random.default_rng(20260713)
boot = {name: [] for name in ["auc", "accuracy", "sensitivity", "specificity", "f1"]}
valid = 0
for _ in range(10000):
    idx = rng.integers(0, len(y), len(y))
    yy = y[idx]
    if yy.min() == yy.max():
        continue
    vals = metric_values(yy, prob[idx], threshold[idx])
    for name in boot:
        boot[name].append(vals[name])
    valid += 1

rows = []
for name in ["auc", "accuracy", "sensitivity", "specificity", "f1"]:
    arr = np.asarray(boot[name], dtype=float)
    rows.append(
        {
            "metric": name,
            "estimate": point[name],
            "ci_low": float(np.nanpercentile(arr, 2.5)),
            "ci_high": float(np.nanpercentile(arr, 97.5)),
            "bootstrap_valid": valid,
            "bootstrap_seed": 20260713,
            "scheme": "A_5fold_x100",
            "threshold_strategy": "nested_inner_oof_youden",
            "pooled_decision": "subject_mean_probability_vs_mean_threshold",
            "n_subjects": len(y),
            "n_positive": int(y.sum()),
            "n_negative": int((1 - y).sum()),
        }
    )
ci = pd.DataFrame(rows)
ci.to_csv(CI_CSV, index=False)
ci_rows = {row["metric"]: row for _, row in ci.iterrows()}

doc = Document(SRC)

summary = (
    "* 대표 성능은 A scheme(5-fold x 100 repeats) subject-level pooled OOF 기준으로 "
    f"AUC {fmt('auc', ci_rows)}, Accuracy {fmt('accuracy', ci_rows)}, "
    f"Sensitivity {fmt('sensitivity', ci_rows)}, Specificity {fmt('specificity', ci_rows)}, "
    f"F1 {fmt('f1', ci_rows)}이다."
)
for paragraph in doc.paragraphs:
    if paragraph.text.startswith("* 대표 성능은 A scheme"):
        paragraph.text = summary
        break

# Table 8: overall scheme comparison, A row includes all metric CIs.
table8 = doc.tables[8]
table8.rows[1].cells[1].text = fmt("auc", ci_rows)
table8.rows[1].cells[2].text = fmt("sensitivity", ci_rows)
table8.rows[1].cells[3].text = fmt("specificity", ci_rows)
table8.rows[1].cells[4].text = fmt("f1", ci_rows)
table8.rows[1].cells[5].text = fmt("accuracy", ci_rows)

# Table 10: nested validation comparison, A row includes all metric CIs.
table10 = doc.tables[10]
table10.rows[1].cells[1].text = fmt("auc", ci_rows)
table10.rows[1].cells[2].text = fmt("sensitivity", ci_rows)
table10.rows[1].cells[3].text = fmt("specificity", ci_rows)
table10.rows[1].cells[4].text = fmt("f1", ci_rows)
table10.rows[1].cells[5].text = fmt("accuracy", ci_rows)

# Table 11: previous-vs-final comparison.
table11 = doc.tables[11]
table11.rows[5].cells[2].text = f"{fmt('auc', ci_rows)}; A pooled OOF"
table11.rows[6].cells[2].text = f"{fmt('sensitivity', ci_rows)}; A pooled OOF"
table11.rows[7].cells[2].text = f"{fmt('specificity', ci_rows)}; A pooled OOF"

note = (
    "Metric 신뢰구간은 A scheme subject-level pooled OOF 평균 확률 및 평균 threshold를 기준으로, "
    "67명을 subject 단위로 10,000회 bootstrap resampling하여 산출하였다. "
    "반복 예측 row가 아니라 subject를 재표본추출 단위로 사용해 동일 대상자의 반복 예측이 CI를 과도하게 좁히지 않도록 했다."
)

inserted_note = False
for i, paragraph in enumerate(doc.paragraphs):
    if "AUC 신뢰구간은 A scheme subject-level pooled OOF 평균 확률" in paragraph.text:
        paragraph.text = note
        inserted_note = True
        break

if not inserted_note:
    for i, paragraph in enumerate(doc.paragraphs):
        if "운동저하 25명 중 20명 정탐" in paragraph.text:
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            doc.paragraphs[i + 1].text = note
            break

doc.save(OUT)
print(OUT)
print(CI_CSV)

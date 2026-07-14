from pathlib import Path

from docx import Document
from docx.shared import Pt


DOCX = Path(r"C:\Users\whdgu\Desktop\파이널 프로젝트\final_보행설명_nestedCV_최종.docx")


def add_paragraph(text="", bold=False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_values in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_values):
            cell.text = text
    doc.add_paragraph()
    return table


doc = Document(DOCX)

doc.add_page_break()
add_paragraph("11. Android SensorManager 수집 및 전처리 명세", bold=True)
add_paragraph(
    "최종 앱에서는 스마트폰을 허리밴드에 세로 방향으로 고정한 상태에서 SensorManager로 IMU를 수집한다. "
    "모델 입력은 원시 센서값 자체가 아니라, 30초 수집 신호 중 품질이 가장 좋은 10초 구간에서 산출한 4개 보행 feature이다. "
    "따라서 앱 구현 시 센서 수집, 축 매핑, 리샘플링, 필터, feature 산출 순서를 학습 파이프라인과 최대한 동일하게 유지해야 한다."
)

add_paragraph("11.1 수집 센서", bold=True)
add_table(
    ["SensorManager 항목", "수집값", "용도", "구현 메모"],
    [
        [
            "TYPE_ACCELEROMETER",
            "x, y, z + timestamp",
            "수직 보행 활력, 좌우 안정성, stride 규칙성 산출",
            "가능하면 SENSOR_DELAY_FASTEST 또는 약 100Hz 수준으로 수집",
        ],
        [
            "TYPE_GYROSCOPE",
            "x, y, z + timestamp",
            "몸통 roll 회전 변동성 산출",
            "roll_amp_pool_iqr 계산에 gyro.z 사용",
        ],
        [
            "TYPE_ROTATION_VECTOR 또는 TYPE_GAME_ROTATION_VECTOR",
            "선택 수집",
            "기기 자세 보정/품질 확인 보조",
            "최종 4개 feature에는 직접 사용하지 않지만, 고급화 시 축 보정에 사용 가능",
        ],
    ],
)

add_paragraph("11.2 축 매핑", bold=True)
add_paragraph(
    "현재 모델 기준 축 매핑은 허리밴드에 스마트폰을 세로로 고정한 조건을 전제로 한다. "
    "부착 방향이 바뀌면 feature 분포가 달라져 모델 성능이 떨어질 수 있으므로, 앱 화면에서 부착 방향을 명확히 안내해야 한다."
)
add_table(
    ["Android 센서 축", "모델 보행 축", "사용 feature"],
    [
        ["acc.z", "v: 수직축", "v_amp_pool_median, base_v_stride_regularity"],
        ["acc.x", "ml: 좌우축", "ml_amp_pool_iqr"],
        ["acc.y", "ap: 전후축", "stride_duration 탐지 보조"],
        ["gyro.z", "roll: 몸통 회전", "roll_amp_pool_iqr"],
    ],
)

add_paragraph("11.3 전처리 순서", bold=True)
add_table(
    ["단계", "처리 내용", "세부 기준"],
    [
        [
            "1. 원시 수집",
            "가속도계와 자이로스코프를 동시에 30초 수집",
            "각 샘플은 timestamp(ns), acc(x,y,z), gyro(x,y,z)를 저장",
        ],
        [
            "2. 시간 정렬",
            "timestamp 기준으로 acc/gyro를 동일 시간축에 정렬",
            "불규칙 샘플링은 보간하여 100Hz 등간격 신호로 리샘플링",
        ],
        [
            "3. 구간 선택",
            "30초 중 보행 품질이 가장 좋은 10초 window 선택",
            "stride regularity가 계산 가능하고 결측/정지 구간이 적은 window 우선",
        ],
        [
            "4. 축 매핑",
            "acc.z -> v, acc.x -> ml, acc.y -> ap, gyro.z -> roll",
            "스마트폰 부착 방향이 바뀌면 이 매핑도 함께 보정 필요",
        ],
        [
            "5. 필터링",
            "Butterworth 4차 필터 적용",
            "v/ml: 0.6-3.0Hz bandpass, roll: 0.5-5.0Hz bandpass, raw acc는 필요 시 20Hz lowpass",
        ],
        [
            "6. feature 산출",
            "10초 window에서 4개 feature 계산",
            "v_amp_pool_median, ml_amp_pool_iqr, base_v_stride_regularity, roll_amp_pool_iqr",
        ],
        [
            "7. 모델 추론",
            "학습된 scaler와 LogisticRegression 계수 적용",
            "probability >= 0.4898이면 운동기능 저하 가능군으로 판정",
        ],
    ],
)

add_paragraph("11.4 최종 모델 입력 feature 계산", bold=True)
add_table(
    ["Feature", "계산 방법", "해석"],
    [
        [
            "v_amp_pool_median",
            "v = bandpass(acc.z, 0.6-3.0Hz) 후 robust abs amplitude의 median",
            "수직 보행 에너지와 상하 움직임 크기",
        ],
        [
            "ml_amp_pool_iqr",
            "ml = bandpass(acc.x, 0.6-3.0Hz) 후 robust abs amplitude의 IQR",
            "좌우 흔들림 변동성",
        ],
        [
            "base_v_stride_regularity",
            "ap 신호 ACF에서 stride lag를 찾고, 해당 lag의 v 신호 ACF 값을 사용",
            "걸음 반복 패턴의 규칙성",
        ],
        [
            "roll_amp_pool_iqr",
            "roll = gyro.z - median(gyro.z) 후 bandpass(0.5-5.0Hz), robust abs amplitude의 IQR",
            "몸통 회전 흔들림 변동성",
        ],
    ],
)

add_paragraph("11.5 앱 구현 시 예외 처리", bold=True)
add_table(
    ["상황", "처리 권장"],
    [
        ["센서 권한 거부 또는 센서 미지원", "측정 불가 안내 후 재시도/수동 평가로 전환"],
        ["실제 sampling rate가 너무 낮음", "100Hz 리샘플링 전 결측이 과도하면 측정 실패 처리"],
        ["10초 window에서 stride peak 탐지 실패", "base_v_stride_regularity 계산 불가로 서비스 불가 또는 재측정"],
        ["휴대폰 부착 방향 오류", "측정 전 안내 화면과 간단한 정지 자세 체크로 방향 확인"],
        ["걷지 않고 정지/흔들기만 한 경우", "신호 품질 점수 낮음으로 재측정 유도"],
    ],
)

add_paragraph(
    "정리하면, SensorManager에서 최종적으로 필요한 필수 원시값은 acc.x/acc.y/acc.z, gyro.x/gyro.y/gyro.z, timestamp이다. "
    "다만 최종 모델에 직접 입력되는 값은 이 원시값이 아니라 동일 전처리 파이프라인을 거쳐 산출한 4개 feature이며, "
    "threshold 0.4898은 nested CV 기준 최종 Youden threshold이다."
)

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.size is None:
            run.font.size = Pt(10.5)

doc.save(DOCX)
print(DOCX)

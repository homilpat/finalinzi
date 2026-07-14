from pathlib import Path

from docx import Document


SRC = Path(r"C:\Users\whdgu\Desktop\final_보행설명.docx")
OUT = Path(r"C:\Users\whdgu\Desktop\파이널 프로젝트\final_보행설명_nestedCV_최종.docx")


def set_para(paragraph, text):
    paragraph.text = text


def set_cell(cell, text):
    cell.text = str(text)


def fill_row(row, values):
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


doc = Document(SRC)

paragraph_updates = {
    3: "2026-07-13",
    36: "4.2 Threshold 결정 방식",
    37: "최종 threshold는 nested CV 구조 안에서 train fold 내부 inner-OOF 예측값으로 Youden index(sensitivity + specificity - 1)를 최대화하는 값으로 결정하였다. 외부 test fold에는 해당 threshold를 독립 적용하여 threshold 선택 과정의 데이터 누수를 방지하였다.",
    51: "모든 CV scheme에서 subject 단위로 train/test를 분리하고, threshold는 train fold 내부 inner-OOF 예측값에서만 결정하였다. 최종 보고 기준은 nested_inner_oof_youden threshold 전략이다.",
    53: "7.1 전체 scheme 비교 (nested_inner_oof_youden threshold)",
    55: "* 대표 성능은 A scheme(5-fold x 100 repeats) subject-level pooled OOF 기준으로 AUC 0.830, sensitivity 0.800, specificity 0.738이다.",
    56: "* LOSO는 test fold = 1명이라 fold-level sensitivity/specificity가 왜곡될 수 있으므로 subject-level pooled 결과만 참고한다.",
    58: "7.2 A scheme pooled 혼동행렬 상세 (nested_inner_oof_youden)",
    60: "운동저하 25명 중 20명 정탐(Sensitivity 0.800) / 정상 42명 중 31명 정탐(Specificity 0.738)",
    62: "7.3 Nested CV 검증 방식별 비교",
    64: "8. 최종 모델 vs 이전 모델 비교",
    72: "model_data = joblib.load('final_motor_domain4_labwalks10_logistic_C0p5_nested_youden.joblib')",
    75: "threshold = model_data[\"threshold\"]  # 0.4898",
    86: "외부 검증 데이터가 없으며, 67명 단일 Labwalks 데이터셋 기반이므로 독립 검증이 필요하다. 따라서 성능은 진단 성능이 아니라 내부 반복 교차검증 기반 선별 성능으로 해석한다.",
}

for idx, text in paragraph_updates.items():
    set_para(doc.paragraphs[idx], text)

# Table 0: overview
fill_row(doc.tables[0].rows[7], ["최종 threshold", "0.4898 (nested_inner_oof_youden)"])

# Table 1: pipeline
fill_row(doc.tables[1].rows[6], ["⑥ Threshold", "nested CV 내부 Youden", "train fold의 inner-OOF 예측값에서 Youden index 최대"])

# Table 4: threshold strategies
threshold_rows = [
    ["전략", "설명", "채택 여부"],
    ["nested_inner_oof_youden", "outer train fold 내부 inner-OOF 예측값에서 Youden index 최대", "채택 (최종 모델)"],
    ["train_sens90_maxspec", "train fold에서 sensitivity >= 0.90 중 specificity 최대", "이전 후보 전략"],
    ["sens80_p20", "train positive 분포 20th percentile", "비교용"],
    ["spec80_p80", "train negative 분포 80th percentile", "비교용"],
]
for row, values in zip(doc.tables[4].rows, threshold_rows):
    fill_row(row, values)

# Table 5: code/file structure, only rows that changed from final model naming.
fill_row(
    doc.tables[5].rows[8],
    [
        "final__2026/03_code/03_validation/RUN_nested_domain4_oof_cv_final2026.py",
        "nested CV + inner-OOF Youden threshold 산출 / subject-level leakage 방지 검증",
    ],
)
fill_row(
    doc.tables[5].rows[10],
    [
        "final__2026/02_model/final_motor_domain4_labwalks10_logistic_C0p5_nested_youden.joblib",
        "최종 배포 모델 파일 (pipeline + 4개 feature + threshold=0.4898 포함)",
    ],
)
fill_row(
    doc.tables[5].rows[11],
    [
        "final__2026/02_model/final_motor_domain4_labwalks10_logistic_C0p5_nested_youden_metadata.json",
        "nested CV 기준 threshold, subject 수, pooled CV 성능, apparent train 성능 메타데이터",
    ],
)

# Table 8: scheme comparison
scheme_rows = [
    ["Scheme", "AUC", "Sensitivity", "Specificity", "F1", "Accuracy", "AUC gap", "Sens gap"],
    ["A: 5-fold x 100 repeats", "0.830", "0.800", "0.738", "0.714", "0.761", "0.034", "0.040"],
    ["B: 3-fold x 100 repeats", "0.827", "0.800", "0.714", "0.702", "0.746", "0.045", "0.080"],
    ["C: 80/20 x 100 repeats", "0.836", "0.760", "0.738", "0.691", "0.746", "0.064", "0.101"],
    ["E: LOSO (pooled)", "0.830", "0.800", "0.690", "0.690", "0.731", "-", "-"],
]
for row, values in zip(doc.tables[8].rows, scheme_rows):
    fill_row(row, values)

# Table 9: confusion matrix, A pooled mean-prob decision.
conf_rows = [
    ["", "예측: 운동저하(1)", "예측: 정상(0)", "합계"],
    ["실제: 운동저하(1)", "TP = 20", "FN = 5", "25명"],
    ["실제: 정상(0)", "FP = 11", "TN = 31", "42명"],
    ["합계", "31명", "36명", "67명"],
]
for row, values in zip(doc.tables[9].rows, conf_rows):
    fill_row(row, values)

# Table 10: validation scheme details under nested Youden.
nested_rows = [
    ["검증 방식", "AUC", "Sensitivity", "Specificity", "F1", "Accuracy", "TP", "FP", "FN", "TN"],
    ["A 5-fold x 100", "0.830", "0.800", "0.738", "0.714", "0.761", "20", "11", "5", "31"],
    ["B 3-fold x 100", "0.827", "0.800", "0.714", "0.702", "0.746", "20", "12", "5", "30"],
    ["C 80/20 x 100", "0.836", "0.760", "0.738", "0.691", "0.746", "19", "11", "6", "31"],
]
for row, values in zip(doc.tables[10].rows, nested_rows):
    fill_row(row, values)

# Table 11: previous vs final comparison.
comparison_updates = {
    5: ["Test AUC", "0.781", "0.830 (A pooled OOF)"],
    6: ["Test Sensitivity", "0.823", "0.800 (A pooled OOF)"],
    7: ["Test Specificity", "0.633", "0.738 (A pooled OOF)"],
    8: ["AUC gap", "0.059", "0.034 (apparent train - pooled test)"],
    9: ["threshold", "0.4223", "0.4898 (nested_inner_oof_youden)"],
    10: ["대상 subjects", "71명", "67명 (CO024, FL020 제외 후 최종 분석 가능 대상)"],
}
for row_idx, values in comparison_updates.items():
    fill_row(doc.tables[11].rows[row_idx], values)

# Table 12: final artifacts.
artifact_rows = [
    ["파일", "내용"],
    [
        "final__2026/02_model/final_motor_domain4_labwalks10_logistic_C0p5_nested_youden.joblib",
        "sklearn Pipeline 포함 최종 모델 파일 / {pipeline, features, threshold=0.4898, threshold_strategy, decision_rule}",
    ],
    [
        "final__2026/02_model/final_motor_domain4_labwalks10_logistic_C0p5_nested_youden_metadata.json",
        "nested CV 기준 threshold, n_subjects, pooled CV 성능, apparent train 성능, 제외 subject 등 메타데이터",
    ],
    [
        "final__2026/02_model/domain4_nested_full_validation_metrics.csv",
        "nested CV 전체 fold별 raw 성능 CSV",
    ],
]
for row, values in zip(doc.tables[12].rows, artifact_rows):
    fill_row(row, values)

doc.save(OUT)
print(OUT)

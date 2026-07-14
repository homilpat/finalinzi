from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt


ROOT = Path(r"C:\Users\whdgu\Desktop\파이널 프로젝트")
SRC = ROOT / "final_보행설명_nestedCV_최종_전체CI포함.docx"
OUT = ROOT / "final_보행설명_nestedCV_최종_전체CI_72h검증포함.docx"
SUMMARY = ROOT / "final__2026" / "05_daily75h_validation" / "daily75h_fixed_model_validation_summary.csv"


def fmt_metric(row, name):
    return f"{row[name]:.3f} ({row[name + '_ci_low']:.3f}-{row[name + '_ci_high']:.3f})"


def add_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for values in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, values):
            cell.text = str(text)
    doc.add_paragraph()
    return table


summary = pd.read_csv(SUMMARY)
doc = Document(SRC)

doc.add_page_break()
add_heading(doc, "12. 72시간 일상 보행 데이터 고정 모델 검증")
doc.add_paragraph(
    "최종 nested CV 모델을 재학습하지 않고 고정한 상태에서, 파이널 보행 프로젝트의 72시간 일상 IMU 데이터에 동일한 10초 window 기반 feature 전처리를 적용해 탐색적 검증을 수행하였다. "
    "단, 대상자가 기존 CO/FL 코호트와 겹치므로 완전한 독립 외부검증이 아니라, 실험실 보행 모델의 free-living/domain-shift 적용 가능성 검토로 해석한다."
)

doc.add_paragraph(
    "라벨은 기존과 동일하게 DGI <= 19 또는 TUG >= 12를 운동기능 저하 가능군(positive=1)으로 정의하였다. "
    "모델 threshold는 nested CV 최종값 0.4898을 그대로 사용하였다."
)

display = summary[summary["cohort"] == "all_matched_valid"].copy()
name_map = {
    "best_window": "best 10초 window",
    "top10_regularity_median": "stride regularity 상위 10% median",
    "all_window_median": "전체 valid window median",
}
rows = []
for _, row in display.iterrows():
    rows.append(
        [
            name_map.get(row["aggregation"], row["aggregation"]),
            f"{int(row['n_subjects'])} ({int(row['n_positive'])}/{int(row['n_negative'])})",
            fmt_metric(row, "auc"),
            fmt_metric(row, "accuracy"),
            fmt_metric(row, "sensitivity"),
            fmt_metric(row, "specificity"),
            fmt_metric(row, "f1"),
            f"TP={int(row['tp'])}, FP={int(row['fp'])}, FN={int(row['fn'])}, TN={int(row['tn'])}",
        ]
    )

add_table(
    doc,
    ["집계 방식", "N (pos/neg)", "AUC (95% CI)", "Accuracy (95% CI)", "Sensitivity (95% CI)", "Specificity (95% CI)", "F1 (95% CI)", "혼동행렬"],
    rows,
)

rep = display[display["aggregation"] == "top10_regularity_median"].iloc[0]
doc.add_paragraph(
    "대표 해석은 stride regularity 상위 10% window의 subject-level median 집계가 가장 균형적이다. "
    f"이 방식에서 AUC {fmt_metric(rep, 'auc')}, sensitivity {fmt_metric(rep, 'sensitivity')}, specificity {fmt_metric(rep, 'specificity')}, "
    f"accuracy {fmt_metric(rep, 'accuracy')}로 나타났다. "
    "best-window 방식은 specificity가 높지만 sensitivity가 낮고, 전체 window median 방식은 sensitivity가 높지만 specificity가 크게 낮아진다."
)

doc.add_paragraph(
    "따라서 72시간 데이터 결과는 모델이 일상 보행에서도 일정 수준의 판별 신호를 유지한다는 근거로 사용할 수 있으나, "
    "성능 저하와 집계 방식 민감성이 확인되므로 정식 외부검증으로 표현하지 않는다. "
    "보고 표현은 '72시간 free-living 데이터에서의 탐색적 도메인 전이 검증'이 적절하다."
)

doc.add_paragraph(
    "산출 파일: final__2026/05_daily75h_validation/daily75h_fixed_model_validation_summary.csv, "
    "daily75h_fixed_model_subject_predictions.csv, daily75h_service10_model_windows_merged.csv"
)

doc.save(OUT)
print(OUT)
